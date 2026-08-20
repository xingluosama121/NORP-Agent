import docx
import sys

doc = docx.Document('whats_new.docx')
with open('extracted_content.txt', 'w', encoding='utf-8') as f:
    f.write(f"Total paragraphs: {len(doc.paragraphs)}\n")
    f.write(f"Total tables: {len(doc.tables)}\n\n")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            f.write(f"[P{i}] [{para.style.name}] {text}\n")
    
    for ti, table in enumerate(doc.tables):
        f.write(f"\n=== TABLE {ti} ===\n")
        for ri, row in enumerate(table.rows):
            cells = [cell.text for cell in row.cells]
            f.write(f"  Row {ri}: {cells}\n")

print("Done extracting")
